import docx
import os
import re
from docx.opc.exceptions import PackageNotFoundError

def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from the DOCX file
    """
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():  # Skip empty cells
                        row_text.append(cell.text.strip())
                if row_text:  # Skip empty rows
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    except PackageNotFoundError:
        raise ValueError(f"Could not open the file at {file_path}. It may not be a valid DOCX file.")
    except Exception as e:
        raise Exception(f"Error processing DOCX file: {str(e)}")

def extract_patent_sections(text):
    """
    Extract standard patent sections from the text.
    
    Args:
        text (str): Full text of the patent document
        
    Returns:
        dict: Dictionary containing different sections of the patent
    """
    sections = {
        "title": "",
        "abstract": "",
        "background": "",
        "summary": "",
        "description": "",
        "claims": [],
        "drawings": []
    }
    
    # Simple pattern matching to identify sections
    # This is a basic implementation and may need adjustments for specific formats
    
    # Find title (usually at the beginning)
    title_match = re.search(r'^(.*?)(?:\n\n|\r\n\r\n)', text, re.DOTALL)
    if title_match:
        sections["title"] = title_match.group(1).strip()
    
    # Find abstract
    abstract_match = re.search(r'(?i)(?:abstract|摘要)\s*[:\n](.*?)(?:\n\n|\r\n\r\n|$)', text, re.DOTALL)
    if abstract_match:
        sections["abstract"] = abstract_match.group(1).strip()
    
    # Find background/field of invention
    background_match = re.search(r'(?i)(?:background|技术背景|发明背景)\s*[:\n](.*?)(?:\n\n|\r\n\r\n|(?i)summary|摘要|$)', text, re.DOTALL)
    if background_match:
        sections["background"] = background_match.group(1).strip()
    
    # Find summary
    summary_match = re.search(r'(?i)(?:summary|发明内容)\s*[:\n](.*?)(?:\n\n|\r\n\r\n|(?i)description|说明书|$)', text, re.DOTALL)
    if summary_match:
        sections["summary"] = summary_match.group(1).strip()
    
    # Find detailed description
    description_match = re.search(r'(?i)(?:detailed description|具体实施方式)\s*[:\n](.*?)(?:\n\n|\r\n\r\n|(?i)claims|权利要求|$)', text, re.DOTALL)
    if description_match:
        sections["description"] = description_match.group(1).strip()
    
    # Find claims
    claims_text = re.search(r'(?i)(?:claims|权利要求)\s*[:\n](.*?)(?:\n\n|\r\n\r\n|(?i)drawings|附图|$)', text, re.DOTALL)
    if claims_text:
        claims_raw = claims_text.group(1).strip()
        # Split claims by numbers (1., 2., etc.)
        claims_list = re.split(r'\n\s*\d+\.', claims_raw)
        # Clean up and add back numbers
        for i, claim in enumerate(claims_list):
            if i == 0 and not re.match(r'^\s*\d+\.', claims_raw):
                # First element is not a claim if the claims don't start with a number
                continue
            sections["claims"].append(claim.strip())
    
    return sections 